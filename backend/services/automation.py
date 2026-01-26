from __future__ import annotations

import argparse
import json
import logging
import os
import re
import sys
import time
import unicodedata
from dataclasses import asdict, dataclass
from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any, Collection, Iterable, Mapping, Sequence, cast
from copy import deepcopy
from threading import Lock

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
import docx.oxml
from PIL import Image

from sqlalchemy import text
from sqlalchemy.engine import Connection
from sqlalchemy.exc import SQLAlchemyError

from backend.core.config import OUTPUT_DIR, database_engine

logger = logging.getLogger("cv_automation")
logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")

# Campos de creditação presentes na tabela de alocação agregada.
ACCREDITATION_COLUMNS = ("AACSB", "EQUIS", "AMBA", "ABET")

PHOTO_TABLE = "faculty_photos"

# Janelas utilizadas para filtrar experiência profissional e produção intelectual.
EXPERIENCE_WINDOW_YEARS = 12
PRODUCTION_WINDOW_YEARS = 5

_YEAR_PATTERN = re.compile(r"(19|20)\d{2}")

# Estruturas de dados 


@dataclass
class EducationRecord:
    """Guarda um curso ou diploma do docente """
    degree: str | None
    institution: str | None
    year: str | None
    country: str | None


@dataclass
class ExperienceEntry:
    """Registra uma experiência profissional relevante """
    role: str | None
    organization: str | None
    city: str | None
    country: str | None
    category: str | None
    start: datetime | None
    end: datetime | None

    def is_within_window(self, years: int, reference: datetime) -> bool:
        if self.start is None and self.end is None:
            return False
        cutoff = reference - timedelta(days=years * 365)
        effective_end = self.end or reference
        return effective_end >= cutoff


@dataclass
class ProductionEntry:
    """Resume uma produção acadêmica que vai para o CV """
    year: str | None
    title: str | None
    production_type: str | None
    nature: str | None
    classification: str | None
    peer_review: str | None
    status_savi: str | None
    status_biblioteca: str | None
    evidence_source: str | None
    lattes_info: str | None


@dataclass
class FacultyProfile:
    """Reúne todas as informações consolidadas de um docente """
    faculty_id: str
    name: str
    email: str | None
    nationality: str | None
    area: str | None
    specialization: str | None
    unit: str | None
    career: str | None
    career_en: str | None
    core_status: str | None
    vertente: str | None
    regime: str | None
    vinculo: str | None
    qualification_summary: str | None
    engagement_description: str | None
    admission_date: datetime | None
    highest_degree: str | None
    time_mission: str | None
    fte: str | None
    teaching_load: str | None
    executive_education_load: str | None
    title_valid_brazil: str | None
    accreditation_flag: str | None
    allocation_tag: str | None
    scholar_profile: str | None
    scopus_profile: str | None
    orcid: str | None
    lattes: str | None
    linkedin: str | None
    personal_site: str | None
    experience_summary: str | None
    international_experience: str | None
    photo: bytes | None
    photo_mime_type: str | None
    photo_filename: str | None
    photo_updated_at: str | None
    education: list[EducationRecord]
    experiences: list[ExperienceEntry]
    productions: list[ProductionEntry]
    phd_title: str | None
    phd_institution: str | None
    phd_year: str | None
    phd_country: str | None
    masters_title: str | None
    masters_institution: str | None
    masters_year: str | None
    masters_country: str | None

    def to_serializable(self) -> dict:
        """Transforma o perfil em um dicionário pronto para JSON """
        return {
            "faculty": {
                "id": self.faculty_id,
                "name": self.name,
                "email": self.email,
                "nationality": self.nationality,
                "area": self.area,
                "specialization": self.specialization,
                "unit": self.unit,
                "career": self.career,
                "career_en": self.career_en,
                "core_status": self.core_status,
                "vertente": self.vertente,
                "regime": self.regime,
                "vinculo": self.vinculo,
                "qualification_summary": self.qualification_summary,
                "engagement_description": self.engagement_description,
                "admission_date": _format_date(self.admission_date),
                "highest_degree": self.highest_degree,
                "time_mission": self.time_mission,
                "fte": self.fte,
                "teaching_load": self.teaching_load,
                "executive_education_load": self.executive_education_load,
                "title_valid_brazil": self.title_valid_brazil,
                "aacsb_flag": self.accreditation_flag,
                "allocation": self.allocation_tag,
                "phd": {
                    "title": self.phd_title,
                    "institution": self.phd_institution,
                    "year": self.phd_year,
                    "country": self.phd_country,
                },
                "masters": {
                    "title": self.masters_title,
                    "institution": self.masters_institution,
                    "year": self.masters_year,
                    "country": self.masters_country,
                },
                "scholar": self.scholar_profile,
                "scopus": self.scopus_profile,
                "orcid": self.orcid,
                "lattes": self.lattes,
                "linkedin": self.linkedin,
                "personal_site": self.personal_site,
                "experience_summary": self.experience_summary,
                "international_experience": self.international_experience,
            },
            "photo": {
                "available": self.photo is not None,
                "mime_type": self.photo_mime_type,
                "filename": self.photo_filename,
                "updated_at": self.photo_updated_at,
            },
            "education": [asdict(record) for record in self.education],
            "experience": [
                {
                    "role": entry.role,
                    "organization": entry.organization,
                    "city": entry.city,
                    "country": entry.country,
                    "category": entry.category,
                    "start": _format_date(entry.start),
                    "end": _format_date(entry.end),
                }
                for entry in self.experiences
            ],
            "production": [
                {
                    "year": entry.year,
                    "title": entry.title,
                    "type": entry.production_type,
                    "nature": entry.nature,
                    "classification": entry.classification,
                    "peer_review": entry.peer_review,
                    "status_savi": entry.status_savi,
                    "status_biblioteca": entry.status_biblioteca,
                    "evidence_source": entry.evidence_source,
                    "lattes_info": entry.lattes_info,
                }
                for entry in self.productions
            ],
        }


def _normalize_token(value: str | None) -> str | None:
    """Normaliza chaves textuais removendo espaços, acentos e padronizando caixa."""
    if value is None:
        return None
    trimmed = " ".join(value.strip().split())
    if not trimmed:
        return None
    decomposed = unicodedata.normalize("NFD", trimmed)
    stripped = "".join(char for char in decomposed if unicodedata.category(char) != "Mn")
    return stripped.upper() or None


def _truthy_flag(value: str | None) -> bool:
    """Converte indicadores textuais em booleano para colunas de creditação."""
    if value is None:
        return False
    normalized = value.strip().upper()
    return normalized in {"SIM", "YES", "TRUE", "1", "Y"}


def _crop_image_to_ratio(image: Image.Image, width_ratio: int, height_ratio: int) -> Image.Image:
    """Faz corte centralizado para adequar a foto ao aspecto exigido pelo template."""
    if width_ratio <= 0 or height_ratio <= 0:
        return image

    width, height = image.size
    if width == 0 or height == 0:
        return image

    target_ratio = width_ratio / height_ratio
    current_ratio = width / height

    if abs(current_ratio - target_ratio) <= 0.001:
        return image

    if current_ratio > target_ratio:
        new_width = int(round(height * target_ratio))
        offset = max((width - new_width) // 2, 0)
        cropped = image.crop((offset, 0, offset + new_width, height))
    else:
        new_height = int(round(width / target_ratio))
        offset = max((height - new_height) // 2, 0)
        cropped = image.crop((0, offset, width, offset + new_height))

    return cropped.copy()


class CVAutomation:
    def __init__(self, output_root: Path = OUTPUT_DIR):
        """Configura caminho do banco e da pasta de saída """
        self.engine = database_engine()
        self.output_root = output_root
        self.output_root.mkdir(parents=True, exist_ok=True)
        ttl_env = os.getenv("AUTOMATION_CACHE_TTL")
        try:
            ttl_value = int(ttl_env) if ttl_env is not None else 120
        except ValueError:
            ttl_value = 120
        self._cache_ttl = max(ttl_value, 0)
        self._cache_lock = Lock()
        self._summary_cache: dict[tuple[Any, ...], tuple[float, list[dict], int]] = {}
        self._profile_cache: dict[str, tuple[float, dict]] = {}
        self._all_profiles_cache: tuple[float, list[dict]] | None = None

    @staticmethod
    def _fetch_all(conn: Connection, sql: str, params: dict | None = None) -> list[dict]:
        rows = conn.execute(text(sql), params or {}).mappings().all()
        return [dict(row) for row in rows]

    @staticmethod
    def _fetch_one(conn: Connection, sql: str, params: dict | None = None) -> dict | None:
        row = conn.execute(text(sql), params or {}).mappings().first()
        return dict(row) if row is not None else None

    @staticmethod
    def _safe_fetch_all(conn: Connection, sql: str, params: dict | None = None) -> list[dict]:
        """Executa consultas opcionalmente falhando em silêncio para consultas auxiliares."""
        try:
            return CVAutomation._fetch_all(conn, sql, params)
        except SQLAlchemyError:
            return []

    def _cache_enabled(self) -> bool:
        """Indica se o cache em memória está habilitado de acordo com o TTL."""
        return self._cache_ttl > 0

    def _cache_now(self) -> float:
        """Retorna um carimbo monotônico usado para validar expiração de cache."""
        return time.monotonic()

    def _cache_expired(self, timestamp: float) -> bool:
        """Determina se um item armazenado já excedeu o TTL configurado."""
        return self._cache_ttl > 0 and (self._cache_now() - timestamp) >= self._cache_ttl

    @staticmethod
    def _cache_clone(value: Any) -> Any:
        """Copia profunda do valor em cache para evitar mutações externas."""
        return deepcopy(value)

    def invalidate_cache(self) -> None:
        """Limpa todas as estruturas de cache para refletir dados atualizados."""
        with self._cache_lock:
            self._summary_cache.clear()
            self._profile_cache.clear()
            self._all_profiles_cache = None

    def _get_summary_cache(self, key: tuple[Any, ...]) -> tuple[list[dict], int] | None:
        """Recupera uma página de resumos do cache quando ainda válida."""
        if not self._cache_enabled():
            return None
        with self._cache_lock:
            entry = self._summary_cache.get(key)
            if not entry:
                return None
            timestamp, data, total = entry
            if self._cache_expired(timestamp):
                del self._summary_cache[key]
                return None
            return self._cache_clone(data), total

    def _set_summary_cache(self, key: tuple[Any, ...], data: list[dict], total: int) -> None:
        """Armazena uma lista de resumos e o total associado para reaproveitamento."""
        if not self._cache_enabled():
            return
        with self._cache_lock:
            self._summary_cache[key] = (self._cache_now(), self._cache_clone(data), total)

    def _get_profile_cache(self, faculty_id: str) -> dict | None:
        """Busca um perfil serializado no cache em memória conforme o TTL."""
        if not self._cache_enabled():
            return None
        with self._cache_lock:
            entry = self._profile_cache.get(faculty_id)
            if not entry:
                return None
            timestamp, payload = entry
            if self._cache_expired(timestamp):
                del self._profile_cache[faculty_id]
                return None
            return self._cache_clone(payload)

    def _set_profile_cache(self, faculty_id: str, payload: dict) -> None:
        """Guarda o perfil serializado de um docente específico."""
        if not self._cache_enabled():
            return
        with self._cache_lock:
            self._profile_cache[faculty_id] = (self._cache_now(), self._cache_clone(payload))

    def _get_all_profiles_cache(self) -> list[dict] | None:
        """Retorna a lista completa de perfis já formatados caso esteja em cache."""
        if not self._cache_enabled():
            return None
        with self._cache_lock:
            entry = self._all_profiles_cache
            if entry is None:
                return None
            timestamp, payload = entry
            if self._cache_expired(timestamp):
                self._all_profiles_cache = None
                return None
            return self._cache_clone(payload)

    def _set_all_profiles_cache(self, payload: list[dict]) -> None:
        """Persiste o cache com todos os perfis para acelerar listagens completas."""
        if not self._cache_enabled():
            return
        with self._cache_lock:
            self._all_profiles_cache = (self._cache_now(), self._cache_clone(payload))

    def run(self, accreditation: str, faculty_ids: Sequence[str] | None = None) -> list[dict]:
        """Cria os arquivos da acreditação pedida """
        accreditation_key = accreditation.strip().upper()
        if not accreditation_key:
            raise ValueError("Accreditation must not be empty")
        planned_ids = list(faculty_ids) if faculty_ids else self._fetch_all_ids()
        logger.info("Starting automation for %s with %d faculty", accreditation_key, len(planned_ids))

        metadata: list[dict] = []
        timestamp = datetime.utcnow().isoformat(timespec="seconds") + "Z"
        accreditation_dir = self.output_root / accreditation_key.lower()
        accreditation_dir.mkdir(parents=True, exist_ok=True)

        with self.engine.connect() as conn:
            for faculty_id in planned_ids:
                # Monta os dados completos do docente direto do banco.
                try:
                    profile = self._build_profile(conn, faculty_id)
                except Exception as exc:  # noqa: BLE001
                    logger.exception("Failed to assemble data for faculty %s: %s", faculty_id, exc)
                    continue

                if profile is None:
                    logger.warning("Skipping faculty %s: no base record found", faculty_id)
                    continue

                limited_profile = profile

                docx_target = accreditation_dir / f"{faculty_id}_{_slugify(profile.name)}.docx"
                json_path = accreditation_dir / f"{faculty_id}_{_slugify(profile.name)}.json"

                try:
                    saved_docx_path = self._generate_document(limited_profile, docx_target, include_photo=True)
                    self._write_json(limited_profile, json_path)
                except Exception as exc:  # noqa: BLE001
                    logger.exception("Failed to render CV for faculty %s: %s", faculty_id, exc)
                    continue

                metadata.append(
                    {
                        "faculty_id": faculty_id,
                        "name": profile.name,
                        "accreditation": accreditation_key,
                        "docx_path": saved_docx_path.relative_to(self.output_root).as_posix(),
                        "json_path": json_path.relative_to(self.output_root).as_posix(),
                        "generated_at": timestamp,
                    }
                )
                logger.info("Generated CV for %s (%s)", profile.name, faculty_id)

        if metadata:
            log_path = accreditation_dir / f"run_{datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')}.json"
            log_path.write_text(json.dumps(metadata, indent=2), encoding="utf-8")
            logger.info("Run summary written to %s", log_path)
        else:
            logger.warning("No CVs generated for %s", accreditation_key)

        return metadata

    def fetch_profile(self, faculty_id: str) -> dict | None:
        """Busca um docente pelo id e devolve os dados prontos para JSON """
        cached = self._get_profile_cache(faculty_id)
        if cached is not None:
            return cached

        with self.engine.connect() as conn:
            try:
                # Reaproveita a lógica interna para compor o perfil.
                profile = self._build_profile(conn, faculty_id)
            except Exception as exc:  # noqa: BLE001
                logger.exception(
                    "Falha ao recuperar dados do docente %s: %s",
                    faculty_id,
                    exc,
                )
                return None

        if profile is None:
            return None

        serialized = profile.to_serializable()
        self._set_profile_cache(faculty_id, serialized)
        return serialized

    def fetch_all_profiles(self) -> list[dict]:
        """Devolve todos os docentes com os dados já formatados """
        cached = self._get_all_profiles_cache()
        if cached is not None:
            return cached

        faculty_ids = self._fetch_all_ids()
        if not faculty_ids:
            return []

        results: list[dict] = []
        with self.engine.connect() as conn:
            for faculty_id in faculty_ids:
                # Continua mesmo se um docente gerar erro
                try:
                    # Evita travar a exportação quando um registro falha
                    profile = self._build_profile(conn, faculty_id)
                except Exception as exc:  # noqa: BLE001
                    logger.exception(
                        "Falha ao recuperar dados do docente %s: %s",
                        faculty_id,
                        exc,
                    )
                    continue

                if profile is None:
                    continue

                results.append(profile.to_serializable())

            self._set_all_profiles_cache(results)
        return results

    def fetch_profiles_summary(
        self,
        *,
        offset: int = 0,
        limit: int = 50,
        allocated_only: bool = True,
        accreditations: Collection[str] | None = None,
    ) -> tuple[list[dict], int]:
        """Lê apenas os campos essenciais para montar os cards rapidamente."""
        if limit <= 0:
            raise ValueError("O parâmetro 'limit' deve ser positivo")

        capped_limit = min(limit, 50)
        safe_offset = max(offset, 0)
        accreditation_filter = {
            item.strip().upper()
            for item in (accreditations or [])
            if isinstance(item, str) and item.strip()
        }
        cache_key = (
            safe_offset,
            capped_limit,
            bool(allocated_only),
            tuple(sorted(accreditation_filter)),
        )

        cached = self._get_summary_cache(cache_key)
        if cached is not None:
            return cached

        with self.engine.connect() as conn:
            base_rows = self._fetch_all(
                conn,
                """
                SELECT id, nome_padrao, area, nova_area, unid_acad
                FROM base_de_dados_docente
                """,
            )

            detail_rows = self._safe_fetch_all(
                conn,
                """
                SELECT nome_completo, disciplina
                FROM alocacao_2026_1_reldetalhe
                """,
            )

            accreditation_rows = self._safe_fetch_all(
                conn,
                """
                SELECT disciplina, aacsb, equis, amba, abet
                FROM alocacao_26_1
                """,
            )

            photo_rows = self._safe_fetch_all(
                conn,
                f"SELECT faculty_id FROM {PHOTO_TABLE}",
            )

        allocation_data_available = bool(detail_rows)
        accreditation_data_available = bool(accreditation_rows)

        # Constrói mapa de disciplinas para os selos de creditação marcados como "SIM".
        discipline_accreditations: dict[str, set[str]] = {}
        for row in accreditation_rows:
            discipline_key = _normalize_token(_row_value(row, "DISCIPLINA", "disciplina"))
            if not discipline_key:
                continue

            flags = {
                column
                for column in ACCREDITATION_COLUMNS
                if _truthy_flag(_row_value(row, column))
            }
            if not flags:
                continue
            discipline_accreditations.setdefault(discipline_key, set()).update(flags)

        # Indexa docentes alocados com contagem e creditações derivadas das disciplinas.
        allocation_index: dict[str, dict[str, Any]] = {}
        for row in detail_rows:
            name_key = _normalize_token(_row_value(row, "nome_completo", "NOME_COMPLETO"))
            if not name_key:
                continue

            entry = allocation_index.setdefault(
                name_key,
                {"count": 0, "accreditations": set()},
            )
            entry["count"] = int(entry["count"]) + 1

            discipline_key = _normalize_token(_row_value(row, "disciplina", "DISCIPLINA"))
            if discipline_key and discipline_key in discipline_accreditations:
                acc_set = cast(set[str], entry["accreditations"])
                acc_set.update(discipline_accreditations[discipline_key])

        photo_ids = {
            str(_row_value(row, "faculty_id", "FACULTY_ID", "id", "ID")).strip()
            for row in photo_rows
            if _row_value(row, "faculty_id", "FACULTY_ID", "id", "ID") is not None
            and str(_row_value(row, "faculty_id", "FACULTY_ID", "id", "ID")).strip()
        }

        summaries: list[dict] = []
        for row in base_rows:
            faculty_id = _row_value(row, "id", "ID")
            name_raw = (str(_row_value(row, "nome_padrao", "NOME_PADRAO")) or "").strip()
            if not faculty_id or not name_raw:
                continue

            name_key = _normalize_token(name_raw)
            allocation_data = allocation_index.get(name_key or "")
            allocation_count = int(allocation_data["count"]) if allocation_data else 0
            accreditation_values = (
                sorted(cast(set[str], allocation_data["accreditations"]))
                if allocation_data
                else []
            )
            accreditation_set = set(accreditation_values)
            has_allocation = allocation_count > 0

            if allocation_data_available and allocated_only and not has_allocation:
                continue

            if accreditation_filter and not accreditation_set.intersection(accreditation_filter):
                continue

            if accreditation_data_available and not accreditation_values:
                continue

            area_code = (str(_row_value(row, "area", "AREA")) or "").strip()
            area_label = _format_area(area_code) if area_code else None
            area_value_raw = _row_value(row, "nova_area", "NOVA_AREA") or area_label
            if isinstance(area_value_raw, str):
                area_value_raw = area_value_raw.strip()
            area_value = area_value_raw or None
            unit_value = (str(_row_value(row, "unid_acad", "UNID_ACAD")) or "").strip() or None

            summaries.append(
                {
                    "id": str(faculty_id),
                    "name": name_raw,
                    "area": area_value,
                    "unit": unit_value,
                    "allocation_count": allocation_count,
                    "accreditations": accreditation_values,
                    "has_allocation": has_allocation,
                    "has_photo": str(faculty_id).strip() in photo_ids,
                }
            )

        summaries.sort(key=lambda item: item["name"].lower())
        total = len(summaries)

        start = min(safe_offset, total)
        end = min(start + capped_limit, total)
        page_slice = summaries[start:end]
        self._set_summary_cache(cache_key, page_slice, total)
        return page_slice, total

    def export_artifact(
        self,
        faculty_id: str,
        export_format: str,
        *,
        include_photo: bool = True,
    ) -> dict | None:
        """Gera um artefato único no formato solicitado."""

        normalized = (export_format or "").strip().lower()
        if not normalized:
            raise ValueError("Formato de exportação é obrigatório")

        if normalized == "docx":
            return self._export_docx(faculty_id, include_photo=include_photo)
        if normalized == "pdf":
            return self._export_pdf(faculty_id, include_photo=include_photo)

        raise ValueError(f"Formato de exportação não suportado: {export_format}")

    def export_doc(self, faculty_id: str, *, include_photo: bool = True) -> dict | None:
        """Compatibilidade retroativa: mantém a exportação em DOCX."""

        return self._export_docx(faculty_id, include_photo=include_photo)

    def _export_docx(self, faculty_id: str, *, include_photo: bool = True) -> dict | None:
        profile = self._load_profile(faculty_id)
        if profile is None:
            return None

        docx_target = self.output_root / f"{faculty_id}_{_slugify(profile.name)}.docx"
        docx_target.parent.mkdir(parents=True, exist_ok=True)
        saved_docx_path = self._generate_document(profile, docx_target, include_photo=include_photo)

        return {
            "faculty_id": faculty_id,
            "name": profile.name,
            "docx_path": saved_docx_path.relative_to(self.output_root).as_posix(),
        }

    def _export_pdf(self, faculty_id: str, *, include_photo: bool = True) -> dict | None:
        # Reutiliza a geração original do DOCX para garantir layout idêntico.
        docx_metadata = self._export_docx(faculty_id, include_photo=include_photo)
        if docx_metadata is None:
            return None

        # O caminho relativo será usado tanto para manter o retorno legado quanto para localizar o arquivo.
        docx_relative = docx_metadata.get("docx_path")
        if not docx_relative:
            return None

        docx_path = (self.output_root / docx_relative).resolve()
        pdf_path = docx_path.with_suffix(".pdf")
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        # Converte o DOCX recém-gerado preservando a formatação original.
        self._convert_docx_to_pdf(docx_path, pdf_path)

        return {
            "faculty_id": docx_metadata.get("faculty_id", faculty_id),
            "name": docx_metadata.get("name"),
            "docx_path": docx_relative,
            "pdf_path": pdf_path.relative_to(self.output_root).as_posix(),
        }

    def _fetch_all_ids(self) -> list[str]:
        with self.engine.connect() as conn:
            ids = conn.execute(text("SELECT id FROM base_de_dados_docente")).scalars().all()
            return [str(identifier) for identifier in ids]

    def _load_profile(self, faculty_id: str) -> FacultyProfile | None:
        with self.engine.connect() as conn:
            return self._build_profile(conn, faculty_id)

    def _build_profile(self, conn: Connection, faculty_id: str) -> FacultyProfile | None:
        """Lê todas as colunas necessárias para gerar um perfil completo """
        faculty_row = self._fetch_one(
            conn,
            """
            SELECT
                id,
                nome_padrao,
                email,
                nacionalidade,
                area,
                nova_area,
                unid_acad,
                carreira,
                carreira_en,
                core_non_core,
                vertente,
                regime,
                v_nculo,
                qualif_descricao_2026_2027,
                engajamento_descricao,
                admissao,
                tit_maxima,
                time_mission,
                fte,
                ch_total_ano_vigente,
                titulo_valido_brasil,
                exp_prof,
                exp_int,
                ch_ed_ex_ano_vigente,
                aacsb_2025,
                aloca_o_2025,
                t_dout_en,
                t_dout_ies,
                t_dout_ano,
                t_dout_pais_en,
                t_mestrado_en,
                t_mestrado_ies,
                t_mestrado_ano,
                t_mestrado_pais_en,
                scholar,
                scopus,
                orcid,
                lattes,
                linkedin,
                site_pessoal
            FROM base_de_dados_docente
            WHERE id = :faculty_id
            """,
            {"faculty_id": faculty_id},
        )

        if faculty_row is None:
            return None

        photo_row = self._fetch_one(
            conn,
            f"SELECT image, mime_type, filename, updated_at FROM {PHOTO_TABLE} WHERE faculty_id = :faculty_id",
            {"faculty_id": faculty_row["id"]},
        )

        photo_blob = photo_row["image"] if photo_row is not None else None
        photo_bytes = bytes(photo_blob) if photo_blob is not None else None
        photo_mime = photo_row["mime_type"] if photo_row is not None else None
        photo_filename = photo_row["filename"] if photo_row is not None else None
        photo_updated_at = photo_row["updated_at"] if photo_row is not None else None

        # Carrega listas auxiliares (experiência, educação, produção).
        reference_point = datetime.utcnow()
        experiences = [
            entry
            for entry in self._load_experience(conn, faculty_row["id"])
            if entry.is_within_window(EXPERIENCE_WINDOW_YEARS, reference_point)
        ]
        education = _build_education_records(faculty_row)
        try:
            raw_productions = list(self._load_production(conn, faculty_row["nome_padrao"]))
        except SQLAlchemyError as exc:
            logger.warning(
                "Skipping production data for %s due to database error: %s",
                faculty_row["nome_padrao"],
                exc,
            )
            raw_productions = []

        productions = [
            entry
            for entry in raw_productions
            if _production_is_recent(entry, reference_point.year, PRODUCTION_WINDOW_YEARS)
        ]

        return FacultyProfile(
            faculty_id=str(faculty_row["id"]),
            name=faculty_row["nome_padrao"],
            email=faculty_row["email"],
            nationality=faculty_row["nacionalidade"],
            area=faculty_row["area"],
            specialization=faculty_row["nova_area"],
            unit=faculty_row["unid_acad"],
            career=faculty_row["carreira"],
            career_en=faculty_row["carreira_en"],
            core_status=faculty_row["core_non_core"],
            vertente=faculty_row["vertente"],
            regime=faculty_row["regime"],
            vinculo=faculty_row["v_nculo"],
            qualification_summary=faculty_row["qualif_descricao_2026_2027"],
            engagement_description=faculty_row["engajamento_descricao"],
            admission_date=_parse_date(faculty_row["admissao"]),
            highest_degree=faculty_row["tit_maxima"],
            time_mission=faculty_row["time_mission"],
            fte=faculty_row["fte"],
            teaching_load=faculty_row["ch_total_ano_vigente"],
            executive_education_load=faculty_row["ch_ed_ex_ano_vigente"],
            title_valid_brazil=faculty_row["titulo_valido_brasil"],
            accreditation_flag=faculty_row["aacsb_2025"],
            allocation_tag=faculty_row["aloca_o_2025"],
            scholar_profile=faculty_row["scholar"],
            scopus_profile=faculty_row["scopus"],
            orcid=faculty_row["orcid"],
            lattes=faculty_row["lattes"],
            linkedin=faculty_row["linkedin"],
            personal_site=faculty_row["site_pessoal"],
            experience_summary=faculty_row["exp_prof"],
            international_experience=faculty_row["exp_int"],
            photo=photo_bytes,
            photo_mime_type=photo_mime,
            photo_filename=photo_filename,
            photo_updated_at=photo_updated_at,
            phd_title=faculty_row["t_dout_en"],
            phd_institution=faculty_row["t_dout_ies"],
            phd_year=faculty_row["t_dout_ano"],
            phd_country=faculty_row["t_dout_pais_en"],
            masters_title=faculty_row["t_mestrado_en"],
            masters_institution=faculty_row["t_mestrado_ies"],
            masters_year=faculty_row["t_mestrado_ano"],
            masters_country=faculty_row["t_mestrado_pais_en"],
            education=education,
            experiences=experiences,
            productions=productions,
        )

    def _load_experience(self, conn: Connection, faculty_id: str) -> Iterable[ExperienceEntry]:
        """Filtra experiências do docente que estejam em inglês """
        rows = conn.execute(
            text(
                """
                SELECT
                    cargo_role,
                    empresa_company,
                    cidade_city,
                    pa_s_country,
                    categoria_prof_res_tch,
                    in_cio,
                    fim,
                    idioma
                FROM docentes_experiencia_profissional
                WHERE id = :faculty_id
                ORDER BY in_cio DESC
                """
            ),
            {"faculty_id": faculty_id},
        ).mappings()

        for row_mapping in rows:
            row = dict(row_mapping)
            language = (row.get("idioma") or "").strip().upper()
            if language and language != "EN":
                continue
            yield ExperienceEntry(
                role=row.get("cargo_role"),
                organization=row.get("empresa_company"),
                city=row.get("cidade_city"),
                country=row.get("pa_s_country"),
                category=row.get("categoria_prof_res_tch"),
                start=_parse_date(row.get("in_cio")),
                end=_parse_date(row.get("fim")),
            )

    def _load_production(self, conn: Connection, faculty_name: str) -> Iterable[ProductionEntry]:
        """Busca produções acadêmicas ordenadas por ano """
        rows = conn.execute(
            text(
                """
                SELECT
                    ano,
                    t_tulo,
                    tipo,
                    ve_culo_ou_natureza,
                    classifica_o,
                    revis_o,
                    status_savi,
                    status_biblioteca,
                    fonte_da_evid_ncia,
                    informa_o_cv_lattes
                FROM docentes_producao
                WHERE professor = :faculty_name
                ORDER BY ano DESC, t_tulo
                """
            ),
            {"faculty_name": faculty_name},
        ).mappings()

        for row_mapping in rows:
            row = dict(row_mapping)
            yield ProductionEntry(
                year=row.get("ano"),
                title=row.get("t_tulo"),
                production_type=row.get("tipo"),
                nature=row.get("ve_culo_ou_natureza"),
                classification=row.get("classifica_o"),
                peer_review=row.get("revis_o"),
                status_savi=row.get("status_savi"),
                status_biblioteca=row.get("status_biblioteca"),
                evidence_source=row.get("fonte_da_evid_ncia"),
                lattes_info=row.get("informa_o_cv_lattes"),
            )

    # ========== Formatação do documento .docx ===========
    def _format_header(self, document, profile: FacultyProfile, *, include_photo: bool = True):
        """Monta cabeçalho com informações básicas do docente."""

        container = document
        paragraph_alignment = 1  # Center by default

        if include_photo and profile.photo:
            table = document.add_table(rows=1, cols=2)
            table.alignment = 0
            table.autofit = False
            table.columns[0].width = Inches(1.25)
            table.columns[1].width = Inches(5.25)

            for row in table.rows:
                for cell in row.cells:
                    cell._element.get_or_add_tcPr().append(
                        docx.oxml.parse_xml(
                            r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                            r'<w:top w:val="none"/>'
                            r'<w:left w:val="none"/>'
                            r'<w:bottom w:val="none"/>'
                            r'<w:right w:val="none"/>'
                            r'</w:tcBorders>'
                        )
                    )

            photo_cell = table.cell(0, 0)
            text_cell = table.cell(0, 1)
            paragraph_alignment = 0  # Left align text when photo is present
            container = text_cell

            text_cell._element.get_or_add_tcPr().append(
                docx.oxml.parse_xml(
                    r'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    r'<w:top w:w="40" w:type="dxa"/>'
                    r'<w:left w:w="40" w:type="dxa"/>'
                    r'<w:bottom w:w="40" w:type="dxa"/>'
                    r'<w:right w:w="28" w:type="dxa"/>'
                    r'</w:tcMar>'
                )
            )

            source_stream = BytesIO(profile.photo)
            source_stream.seek(0)
            image = Image.open(source_stream)
            image.load()
            cropped_image = _crop_image_to_ratio(image, 3, 4)
            output_stream = BytesIO()
            image_format = (image.format or "PNG").upper()
            if image_format not in {"JPEG", "JPG", "PNG"}:
                image_format = "PNG"
            if image_format in {"JPEG", "JPG"} and cropped_image.mode in {"RGBA", "LA"}:
                cropped_image = cropped_image.convert("RGB")
            cropped_image.save(output_stream, format=image_format)
            image.close()
            source_stream.close()
            output_stream.seek(0)

            photo_paragraph = photo_cell.paragraphs[0] if photo_cell.paragraphs else photo_cell.add_paragraph()
            for run in list(photo_paragraph.runs):
                photo_paragraph._element.remove(run._element)
            run = photo_paragraph.add_run()
            run.add_picture(output_stream, width=Inches(1.35))
            photo_paragraph.alignment = 0
            photo_paragraph.paragraph_format.space_after = Pt(0)

        def add_header_line(text: str | None, *, bold: bool = True, size: int = 12, space_after: int = 2) -> None:
            if not text:
                return

            if isinstance(container, DocxDocument):
                paragraph = container.add_paragraph()
            else:
                if container.paragraphs:
                    paragraph = container.paragraphs[0]
                    if paragraph.text:
                        paragraph = container.add_paragraph()
                    else:
                        for run in list(paragraph.runs):
                            paragraph._element.remove(run._element)
                else:
                    paragraph = container.add_paragraph()

            run = paragraph.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(size)
            run.font.bold = bold
            paragraph.alignment = paragraph_alignment
            paragraph.paragraph_format.space_after = Pt(space_after)

        add_header_line(profile.name, size=14, space_after=3)
        add_header_line("Insper Instituto de Ensino e Pesquisa")
        add_header_line(profile.career_en or profile.career)
        add_header_line(_format_area(profile.area or ""))
        add_header_line(profile.email)

        if profile.admission_date:
            formatted_date = _format_date(profile.admission_date)
            add_header_line(f"Admission: {formatted_date}")

        if profile.lattes:
            add_header_line(profile.lattes, bold=False, space_after=12)
        
        # Academic Unit and Nationality
        dict_units={
            'M&E': "Management and Economics",
            'LAW': "Law",
            'ENG&CC': "Engineering and Computer Science",
        }
        unit_text = f"Academic Unit: {dict_units.get(profile.unit, profile.unit)}"
        nationality_text = f"Nationality: {profile.nationality}"

        # Specialization/Area e Nationality na mesma linha
        if profile.unit or profile.nationality:
            # Criar tabela com 1 linha e 2 colunas
            table = document.add_table(rows=1, cols=2)
            table.alignment = 0  # Alinhamento à esquerda
            
            # Remover bordas da tabela
            for row in table.rows:
                for cell in row.cells:
                    cell._element.get_or_add_tcPr().append(
                        docx.oxml.parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
                    )
            
            # Célula esquerda - Especialização
            left_cell = table.rows[0].cells[0]
            if profile.specialization:
                left_para = left_cell.paragraphs[0]
                run_left = left_para.add_run(unit_text)
                run_left.font.name = 'Times New Roman'
                run_left.font.size = Pt(12)
            
            # Célula direita - Nacionalidade
            right_cell = table.rows[0].cells[1]
            if profile.nationality:
                right_para = right_cell.paragraphs[0]
                right_para.alignment = 2
                run_right = right_para.add_run(nationality_text)
                run_right.font.name = 'Times New Roman'
                run_right.font.size = Pt(12)
            
            # Espaçamento após a tabela
            table.rows[0].height = Pt(14)

    # =========================================================================
    # MÉTODOS AUXILIARES DE FORMATAÇÃO E LAYOUT
    # =========================================================================

    def _add_section_header(self, document, text: str):
        """Adiciona um cabeçalho H1 com a formatação específica (Borda inferior) """
        heading = document.add_heading(text, level=1)
        
        # Formatação do Texto
        if heading.runs:
            run = heading.runs[0]
        else:
            run = heading.add_run()
            
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), 'Times New Roman')
        
        # Espaçamento
        heading.paragraph_format.space_after = Pt(6)
        heading.paragraph_format.space_before = Pt(12)
        
        # Borda Inferior (XML Injection)
        p_pr = heading._element.get_or_add_pPr()
        p_borders = docx.oxml.parse_xml(
            r'<w:pBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            r'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>'
            r'</w:pBorders>'
        )
        p_pr.append(p_borders)
        return heading

    def _add_subheader(self, document, text: str):
        """Adiciona subtítulo simples (apenas negrito, sem borda) """
        p = document.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(6)

    def _create_layout_table(self, document):
        """Cria a tabela base de 2 colunas para layout de data/conteúdo """
        table = document.add_table(rows=0, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(0.85)
        table.columns[1].width = Inches(5.65)
        return table

    def _add_layout_row(self, table, left_text: str, right_text: str):
        """Adiciona uma linha formatada à tabela e remove as bordas """
        row = table.add_row()
        
        # Coluna 1: Data/Ano
        row.cells[0].text = left_text
        if row.cells[0].paragraphs[0].runs:
            run = row.cells[0].paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

        # Coluna 2: Conteúdo
        row.cells[1].text = right_text
        if row.cells[1].paragraphs[0].runs:
            run = row.cells[1].paragraphs[0].runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

        # Remove bordas da célula (XML Injection)
        for cell in row.cells:
            tc_pr = cell._element.get_or_add_tcPr()
            tc_borders = docx.oxml.parse_xml(
                r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:top w:val="none"/>'
                r'<w:left w:val="none"/>'
                r'<w:bottom w:val="none"/>'
                r'<w:right w:val="none"/>'
                r'</w:tcBorders>'
            )
            tc_pr.append(tc_borders)

    def _generate_document(
        self,
        profile: FacultyProfile,
        destination: Path,
        *,
        include_photo: bool = True,
    ) -> Path:
        
        document = Document()  # Cria documento vazio para montar o CV.
        
        # Ajuste global de margens se necessário
        # section = document.sections[0]
        # section.left_margin = Inches(1.0)
        # section.right_margin = Inches(1.0)

        self._format_header(document, profile, include_photo=include_photo)

        # ========== EDUCATION ==========
        
        if profile.education:
            self._add_section_header(document, "EDUCATION")
            
            table = self._create_layout_table(document)
            table.columns[0].width = Inches(0.7)  # Mantendo a largura específica usada anteriormente

            for record in profile.education:
                year_text = str(record.year).strip() if record.year else ""
                
                parts = []
                if record.degree: parts.append(record.degree)
                if record.institution: parts.append(record.institution)
                if record.country: parts.append(record.country)
                full_text = ", ".join(parts)
                
                self._add_layout_row(table, year_text, full_text)

            document.add_paragraph().paragraph_format.space_after = Pt(6)
            
            
        # ========== PROFESSIONAL EXPERIENCE ==========
        
        prof_exps = [e for e in profile.experiences if e.category == 'Professional']
        
        if prof_exps:
            self._add_section_header(document, "PROFESSIONAL EXPERIENCE")
            self._add_subheader(document, "Professional experience")
            
            table = self._create_layout_table(document)
            
            for exp in prof_exps:
                date_str = ""
                if exp.start:
                    year_start = exp.start.year
                    if not exp.end:
                        date_str = f"Since {year_start}"
                    else:
                        year_end = exp.end.year
                        date_str = f"{year_start}-{year_end}"
                
                parts = [p for p in [exp.role, exp.organization] if p]
                desc_str = " - ".join(parts)
                if exp.country:
                    desc_str += f", {exp.country}"
                
                self._add_layout_row(table, date_str, desc_str)
            
            document.add_paragraph().paragraph_format.space_after = Pt(6)


        # ========== RESEARCH ACTIVITIES ==========
        
        res_exps = [e for e in profile.experiences if e.category == 'Research']
        
        if res_exps:
            self._add_section_header(document, "RESEARCH ACTIVITIES")
            self._add_subheader(document, "Research Activities & Institutional Contribution")
            
            table = self._create_layout_table(document)
            
            for exp in res_exps:
                date_str = ""
                if exp.start:
                    if not exp.end:
                        date_str = f"Since {exp.start.year}"
                    else:
                        date_str = f"{exp.start.year}-{exp.end.year}"
                
                desc_str = f"{exp.role} - {exp.organization}"
                self._add_layout_row(table, date_str, desc_str)

            document.add_paragraph().paragraph_format.space_after = Pt(6)


        # ========== TEACHING EXPERIENCE ==========
        acad_exps = [e for e in profile.experiences if e.category == 'Academic']
        
        if acad_exps:
            self._add_section_header(document, "TEACHING EXPERIENCE")
            self._add_subheader(document, "Teaching Experience")
            
            table = self._create_layout_table(document)
            
            for exp in acad_exps:
                date_str = ""
                if exp.start:
                    if not exp.end:
                        date_str = f"Since {exp.start.year}"
                    else:
                        date_str = f"{exp.start.year}-{exp.end.year}"
                
                parts = [p for p in [exp.role, exp.organization] if p]
                desc_str = ", ".join(parts)
                if exp.country:
                    desc_str += f", {exp.country}"
                    
                self._add_layout_row(table, date_str, desc_str)
            
            document.add_paragraph().paragraph_format.space_after = Pt(6)


        # ========== INTELLECTUAL CONTRIBUTIONS ==========

        if profile.productions:
            self._add_section_header(document, "INTELLECTUAL CONTRIBUTIONS")
            
            grouped_prod = {}
            for prod in profile.productions:
                p_type = prod.production_type or "Other Productions"
                # Ajuste de categorias baseado em palavras-chave comuns
                if "Artigos" in p_type or "journal" in (prod.nature or "").lower():
                    p_type = "Peer-reviewed Articles"
                elif "Capítulo" in p_type or "Chapter" in p_type:
                    p_type = "Book Chapters"
                
                if p_type not in grouped_prod:
                    grouped_prod[p_type] = []
                grouped_prod[p_type].append(prod)

            priority_order = ["Peer-reviewed Articles", "Book Chapters", "Books", "Other Productions"]
            sorted_keys = sorted(grouped_prod.keys(), key=lambda k: priority_order.index(k) if k in priority_order else 99)

            for p_type in sorted_keys:
                self._add_subheader(document, p_type)
                
                items = grouped_prod[p_type]
                # Ordenar por ano decrescente
                items.sort(key=lambda x: x.year or "0", reverse=True)
                
                for idx, item in enumerate(items, 1):
                    p = document.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    
                    # Número
                    run_num = p.add_run(f"{idx}. ")
                    run_num.font.name = 'Times New Roman'
                    run_num.font.size = Pt(11)
                    
                    # Título
                    run_title = p.add_run(f"{item.title or ''} ")
                    run_title.font.name = 'Times New Roman'
                    run_title.font.size = Pt(11)
                    
                    # (Ano)
                    if item.year:
                        run_year = p.add_run(f"({item.year}). ")
                        run_year.font.name = 'Times New Roman'
                        run_year.font.size = Pt(11)
                    
                    # Natureza / Journal
                    if item.nature:
                        run_nature = p.add_run(f"{item.nature} ")
                        run_nature.font.name = 'Times New Roman'
                        run_nature.font.size = Pt(11)
                        run_nature.font.italic = True

        # Metadados do documento
        document.core_properties.author = "CV Automation"
        document.core_properties.subject = f"{profile.name} CV"

        buffer = BytesIO()
        document.save(buffer)
        content = buffer.getvalue()
        buffer.close()

        destination.parent.mkdir(parents=True, exist_ok=True)

        try:
            with destination.open("wb") as handle:
                handle.write(content)
            saved_path = destination
        except PermissionError:
            timestamp = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
            fallback = destination.with_name(f"{destination.stem}_{timestamp}{destination.suffix}")
            with fallback.open("wb") as handle:
                handle.write(content)
            logger.warning(
                "Permission denied writing %s. Saved document to %s instead.",
                destination,
                fallback,
            )
            saved_path = fallback

        return saved_path

    def _convert_docx_to_pdf(self, source: Path, destination: Path) -> None:
        """Converte um DOCX já formatado em PDF reaproveitando o layout existente."""

        if not source.exists():
            raise ValueError(f"Arquivo DOCX {source} inexistente para conversão em PDF")

        if sys.platform != "win32":
            raise ValueError("Conversão para PDF requer Windows com Microsoft Word instalado")

        try:
            import win32com.client  # type: ignore
        except ImportError as exc:  # pragma: no cover - dependência faltante
            raise ValueError("Pacote pywin32 é necessário para gerar PDF") from exc

        pythoncom = None
        try:
            import pythoncom  # type: ignore
            pythoncom.CoInitialize()
        except ImportError:
            pythoncom = None
        except Exception as exc:  # noqa: BLE001
            logger.warning("Falha ao inicializar pythoncom: %s", exc)

        last_error: Exception | None = None
        try:
            for attempt in range(3):
                word = None
                doc = None
                try:
                    word = win32com.client.DispatchEx("Word.Application")
                    word.Visible = False
                    word.DisplayAlerts = 0

                    # ReadOnly evita conflitos de arquivo enquanto o Word prepara o PDF novo.
                    doc = word.Documents.Open(str(source), ReadOnly=True)
                    doc.SaveAs(str(destination), FileFormat=17)
                    doc.Close(False)
                    return
                except Exception as exc:  # noqa: BLE001
                    last_error = exc
                    logger.warning(
                        "Tentativa %d de converter %s->%s falhou: %s",
                        attempt + 1,
                        source,
                        destination,
                        exc,
                    )
                finally:
                    if doc is not None:
                        try:
                            doc.Close(False)
                        except Exception:  # noqa: BLE001
                            pass
                    if word is not None:
                        try:
                            word.Quit()
                        except Exception:  # noqa: BLE001
                            pass

                if attempt < 2:
                    time.sleep(1.5 * (attempt + 1))

        finally:
            if pythoncom is not None:
                try:
                    pythoncom.CoUninitialize()
                except Exception:  # noqa: BLE001
                    pass

        if last_error is None:
            raise ValueError("Falha desconhecida ao converter DOCX em PDF")
        raise ValueError(f"Falha ao converter DOCX em PDF: {last_error}") from last_error

    def _write_json(self, profile: FacultyProfile, destination: Path) -> None:
        """Serializa o perfil completo em JSON para facilitar integrações externas."""
        destination.write_text(json.dumps(profile.to_serializable(), indent=2), encoding="utf-8")


def _append_table_rows(table, rows: list[tuple[str, str | None]]) -> None:
    """Ajuda a preencher tabelas simples sem repetir código """
    for index, (label, value) in enumerate(rows):
        if index == 0 and len(table.rows) == 1 and not table.rows[0].cells[0].text:
            row = table.rows[0]
        else:
            row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = str(value).strip() if value and str(value).strip() else ""


def _row_value(row: Mapping[str, Any], *keys: str) -> Any:
    """Obtém valores de dicionários oriundos do banco lidando com colunas renomeadas."""
    for key in keys:
        if key in row:
            return row[key]
    normalized_candidates = []
    for key in keys:
        lower = key.lower()
        normalized_candidates.append(lower)
        normalized_candidates.append(lower.replace(" ", "_"))
        normalized_candidates.append(lower.replace(" ", "_").replace("-", "_"))
    for candidate in normalized_candidates:
        if candidate in row:
            return row[candidate]
    return None


def _build_education_records(row: Mapping[str, Any]) -> list[EducationRecord]:
    """Monta lista com as titulações mais relevantes do docente """
    education: list[EducationRecord] = []
    if row["t_dout_en"]:
        education.append(
            EducationRecord(
                degree=row["t_dout_en"],
                institution=row["t_dout_ies"],
                year=row["t_dout_ano"],
                country=row["t_dout_pais_en"],
            )
        )
    if row["t_mestrado_en"]:
        education.append(
            EducationRecord(
                degree=row["t_mestrado_en"],
                institution=row["t_mestrado_ies"],
                year=row["t_mestrado_ano"],
                country=row["t_mestrado_pais_en"],
            )
        )
    if row["tit_maxima"] and not education:
        education.append(
            EducationRecord(
                degree=row["tit_maxima"],
                institution=None,
                year=None,
                country=None,
            )
        )
    return education


def _production_is_recent(entry: ProductionEntry, reference_year: int, window_years: int) -> bool:
    """Indica se a produção está dentro da janela de anos configurada."""
    extracted_year = _extract_year(entry.year)
    if extracted_year is None:
        return False
    cutoff_year = reference_year - window_years + 1
    return extracted_year >= cutoff_year


def _extract_year(value: str | None) -> int | None:
    """Tenta localizar um ano de quatro dígitos no texto informado."""
    if not value:
        return None
    match = _YEAR_PATTERN.search(str(value))
    if not match:
        return None
    try:
        return int(match.group(0))
    except ValueError:
        return None


def _parse_date(raw: str | None) -> datetime | None:
    """Converte textos de data em objetos datetime tratados """
    if not raw:
        return None
    text = raw.strip()
    if not text or text.upper() in {"SEM INFORMAÇÃO", "SEM INFORMACAO", "NSA"}:
        return None

    for fmt in ("%d/%m/%Y", "%m/%d/%Y", "%Y"):
        try:
            return datetime.strptime(text, fmt)
        except ValueError:
            continue
    return None


def _format_date(value: datetime | None) -> str | None:
    """Formata datas no padrão "Mes Ano" esperado pelo relatório """
    if value is None:
        return None
    return value.strftime("%b %Y")


def _slugify(text: str) -> str:
    """Gera um identificador seguro para nomes de arquivo """
    safe = [c.lower() if c.isalnum() else "-" for c in text]
    slug = "".join(safe).strip("-")
    while "--" in slug:
        slug = slug.replace("--", "-")
    return slug or "faculty"


# Mapeia códigos históricos de área para os rótulos exibidos no resumo.
dict_areas = {
    "FIN": "Finance",
    "MGT": "Management",
    "QTM": "Quantitative Methods",
    "NSA": "No Specific Area",
    "LEG": "Legal Studies",
    "ECO": "Economics",
    "MKT": "Marketing",
    "ACC": "Accounting",
    "ITO": "IT and Operations",
}

def _format_area(text: str) -> str:
    """Traduz códigos de área para descrições amigáveis """
    safe = ""
    for c in text.strip():
        safe += c.upper()
    return dict_areas.get(safe, text)


def main() -> None:
    """Ponto de entrada para execução via CLI, espelhando o script original."""
    parser = argparse.ArgumentParser(description="Generate summarized accreditation CVs")
    parser.add_argument("--accreditation", required=True)
    parser.add_argument("--faculty", nargs="*", help="Optional list of faculty IDs to process")
    args = parser.parse_args()

    automation = CVAutomation()
    result = automation.run(args.accreditation, args.faculty)
    logger.info("Completed automation with %d CVs", len(result))


if __name__ == "__main__":
    main()